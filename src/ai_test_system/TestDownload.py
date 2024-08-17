from flask_cors import CORS
from flask import Flask, request, jsonify, Response, send_file
import random
import regex
import os
import sys
import random
from utils import Logger
from spark.SparkApi import SparkLLM
from audio_to_txt.Ifasr_app import audio2txt_Api
from database.Database import get_user_tests_analysis
import json5
from docx import Document
from App import app

appid = 'e76d7d8f'
api_secret = 'Y2Y2ODc2OGQyOWFjMWZhY2JkOTllMDVl'
api_key = '990e2770b030441fbcc126c691daf5cd'

domain = "generalv3.5"    # v3.0版本
Spark_url = "wss://spark-api.xf-yun.com/v3.5/chat"  # v3.5环服务地址

# 创建SparkLLM对象
llm = SparkLLM(appid, api_key, api_secret, Spark_url, domain, False)

get_problems_prompt = '''我以以下格式:
        {
                "subjects": "数学计算" /*问题考察的知识点*/, 
            "count": 3 /*要求提供的题目数量, 不要少于此数量*/,
            "min_difficulty": 3 /*1-10*/,
            "max_difficulty": 8 /*1-10*/,
            "type": ["单选题", "填空题", "判断题"], /*需要的题目类型, 不要给出这里所提到的类型这之外的题目*/
                "others": "希望能出一些计算量比较大的题目"
        }
        给你提供一个对于试题描述的json，请你根据里面的信息以我指定的格式:
        {
                "problems": [
                    {
                        "type": "single_choice" /*注意注意type只能为single_choice, judgement或fillin, 不要给出这之外的类型*/,
                        "problem": "1+1=( )",
                        "choices": ["1","2","3","4"],
                        "answer": [1] /*对应上一个的下标, 通过这里判断单选多选*/,
                        "analysis": "一个很简单的数学题"
                    },
                    {
                        "type": "fillin" /*填空、主观等填写的题目*/,
                        "problem": ["1+2=","4+5=","请回答"], /*以空为分隔符分隔，最后即时没字符也应该有一个结束字符串*/
                        "answer": ["3", "9"] /*几个空几个答案*/,
                        "analysis": "可以化为2进制去计算"
                    },
                    {
                        "type": "judgement",
                        "problem": "6+7=11",
                        "answer": false,
                        "analysis": "可以按计算器"
                    }
                ]
            }
        返回其要求的试题数据。注意只返回5道单选，5道判断，5道填空题。但如果试题描述的"others"要求了题型和题目数量则按照"others"生成题目。
        提供给你的json如下, 请直接返回json, 不要添加任何其他描述:'''

get_problems_prompt1 = '''请以我指定的格式:
        {
                "problems": [
                    {
                        "type": "single_choice" /*注意注意type只能为single_choice, judgement或fillin, 不要给出这之外的类型*/,
                        "problem": "1+1=( )",
                        "choices": ["1","2","3","4"],
                        "answer": [1] /*对应上一个的下标, 通过这里判断单选多选*/,
                        "analysis": "一个很简单的数学题"
                    },
                    {
                        "type": "fillin" /*填空、主观等填写的题目*/,
                        "problem": ["1+2=","4+5=","请回答"], /*以空为分隔符分隔，最后即时没字符也应该有一个结束字符串*/
                        "answer": ["3", "9"] /*几个空几个答案*/,
                        "analysis": "可以化为2进制去计算"
                    },
                    {
                        "type": "judgement",
                        "problem": "6+7=11",
                        "answer": false,
                        "analysis": "可以按计算器"
                    }
                ]
            }
        返回要求的试题数据。注意只返回5道单选，5道判断，5道填空题。请直接返回json, 不要添加任何其他描述'''

@app.route('/get_downloadproblems_txt', methods=['POST'])
def get_downloadproblems_txt_handler():
    content = request.json
    Logger.info(content)
    """
    {
    "subjects": ["math","chinese"],
   	"time": 10 /*mins*/,
   	"min_difficulty": 3 /*1-10*/,
   	"max_difficulty": 8,
   	"type": ["single_choice", "judgement"],
    "others": "希望能出一些计算量比较大的题目"
   }
    """
    subjects = content['subjects']
    time = content['time']
    types = content['type']
    str_types = ['"'+x+'"' for x in types]
    cnt = max(len(subjects), int(time / 3))
    rest = len(subjects)
    ans = {'problems': []}
    for subject in subjects:
        rest -= 1
        now_cnt = random.randint(1, cnt - rest)
        cnt -= now_cnt
        json_str = f'''{{
        "subjects": "{subject}" /*问题考察的知识点*/,
        "count": {now_cnt} /*要求提供的题目数量, 不要少于此数量或多于此数量*/,
        "min_difficulty": {content['min_difficulty']} /*1-10*/,
        "max_difficulty": {content['max_difficulty']} /*1-10*/,
        "type": {str_types} /*需要的题目类型, 不要给出这里所提到的类型这之外的题目*/, 
        "others": "{content['others']}" /*其他要求*/
        }}'''
        Logger.info(f'json_str:{json_str}')
        question = get_problems_prompt + json_str
        ans_str = llm.query(question)
        ans_str = ans_str[ans_str.find('{'): ans_str.rfind('}') + 1]
        # Logger.info(f'ans_str:{ans_str}')
        now_ans = json5.loads(ans_str)
        ans['problems'].extend(now_ans['problems'])
    paper_text = generate_paper_text(ans)
    filepath = os.path.join('ai_test_system','paper')
    os.makedirs(filepath, exist_ok=True)
    filename = os.path.join(filepath,'test.txt')
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(paper_text)
    # 创建下载文件的 Response
    response = Response(paper_text, content_type='text/plain')
    response.headers["Content-Disposition"] = "attachment; filename=problems.txt"
    return response

@app.route('/get_downloadproblems_docx', methods=['POST'])
def get_downloadproblems_docx_handler():
    content = request.json
    # Logger.info(content)
    """
    {
    "subjects": ["math","chinese"],
   	"time": 10 /*mins*/,
   	"min_difficulty": 3 /*1-10*/,
   	"max_difficulty": 8,
   	"type": ["single_choice", "judgement"],
    "others": "希望能出一些计算量比较大的题目"
   }
    """
    subjects = content['subjects']
    time = content['time']
    types = content['type']
    str_types = ['"'+x+'"' for x in types]
    cnt = max(len(subjects), int(time / 3))
    rest = len(subjects)
    ans = {'problems': []}
    for subject in subjects:
        rest -= 1
        now_cnt = random.randint(1, cnt - rest)
        cnt -= now_cnt
        json_str = f'''{{
        "subjects": "{subject}" /*问题考察的知识点*/,
        "count": {now_cnt} /*要求提供的题目数量, 不要少于此数量或多于此数量*/,
        "min_difficulty": {content['min_difficulty']} /*1-10*/,
        "max_difficulty": {content['max_difficulty']} /*1-10*/,
        "type": {str_types} /*需要的题目类型, 不要给出这里所提到的类型这之外的题目*/, 
        "others": "{content['others']}" /*其他要求*/
        }}'''
        Logger.info(f'json_str:{json_str}')
        question = get_problems_prompt + json_str
        ans_str = llm.query(question)
        ans_str = ans_str[ans_str.find('{'): ans_str.rfind('}') + 1]
        # Logger.info(f'ans_str:{ans_str}')
        now_ans = json5.loads(ans_str)
        ans['problems'].extend(now_ans['problems'])
    paper_text = generate_paper_text(ans)
    # 创建 Word 文档
    document = Document()
    document.add_heading('测试试卷', 0)
    document.add_paragraph(paper_text)

    # 保存文档
    filepath = os.path.join('ai_test_system', 'paper')
    os.makedirs(filepath, exist_ok=True)
    filename = os.path.join(filepath, 'test.docx')
    document.save(filename)

    # 返回 Word 文档供下载
    return send_file(filename, as_attachment=True, download_name='problems.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/get_download_customproblems_txt', methods=['POST'])
def get_download_customproblems_txt_handler():
    content = request.json
    Logger.info(content)
    username = content['username']
    subjects = content['subjects']
    time = content['time']
    types = content['type']
    str_types = ['"'+x+'"' for x in types]
    cnt = max(len(subjects), int(time / 3))
    rest = len(subjects)
    user_tests_data = get_user_tests_analysis(username)
    print("user_tests_data:",str(user_tests_data),type(str(user_tests_data)))
    ans = {'problems': []}
    for subject in subjects:
        rest -= 1
        now_cnt = random.randint(1, cnt - rest)
        cnt -= now_cnt
        json_str = f'''{{
        "subjects": "{subject}" /*问题考察的知识点*/,
        "count": {now_cnt} /*要求提供的题目数量, 不要少于此数量或多于此数量*/,
        "min_difficulty": {content['min_difficulty']} /*1-10*/,
        "max_difficulty": {content['max_difficulty']} /*1-10*/,
        "type": {str_types} /*需要的题目类型, 不要给出这里所提到的类型这之外的题目*/, 
        "others": "{content['others']}" /*其他要求*/
        }}'''
        # Logger.info(f'json_str:{json_str}')
        question = get_problems_prompt + json_str + "。同时我提供用户的测试历史数据，请根据用户测试历史微调不同题目的难度：" + str(user_tests_data)
        print("question: ",question)
        ans_str = llm.query(question)
        ans_str = ans_str[ans_str.find('{'): ans_str.rfind('}') + 1]
        # Logger.info(f'ans_str:{ans_str}')
        now_ans = json5.loads(ans_str)
        ans['problems'].extend(now_ans['problems'])
    paper_text = generate_paper_text(ans)
    filepath = os.path.join('ai_test_system','paper')
    os.makedirs(filepath, exist_ok=True)
    filename = os.path.join(filepath,'test.txt')
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(paper_text)
    # 创建下载文件的 Response
    response = Response(paper_text, content_type='text/plain')
    response.headers["Content-Disposition"] = "attachment; filename=problems.txt"
    return response

@app.route('/get_download_customproblems_docx', methods=['POST'])
def get_download_customproblems_docx_handler():
    content = request.json
    Logger.info(content)
    username = content['username']
    subjects = content['subjects']
    time = content['time']
    types = content['type']
    str_types = ['"'+x+'"' for x in types]
    cnt = max(len(subjects), int(time / 3))
    rest = len(subjects)
    user_tests_data = get_user_tests_analysis(username)
    print("user_tests_data:",str(user_tests_data),type(str(user_tests_data)))
    ans = {'problems': []}
    for subject in subjects:
        rest -= 1
        now_cnt = random.randint(1, cnt - rest)
        cnt -= now_cnt
        json_str = f'''{{
        "subjects": "{subject}" /*问题考察的知识点*/,
        "count": {now_cnt} /*要求提供的题目数量, 不要少于此数量或多于此数量*/,
        "min_difficulty": {content['min_difficulty']} /*1-10*/,
        "max_difficulty": {content['max_difficulty']} /*1-10*/,
        "type": {str_types} /*需要的题目类型, 不要给出这里所提到的类型这之外的题目*/, 
        "others": "{content['others']}" /*其他要求*/
        }}'''
        # Logger.info(f'json_str:{json_str}')
        question = get_problems_prompt + json_str + "。同时我提供用户的测试历史数据，请根据用户测试历史微调不同题目的难度：" + str(user_tests_data)
        print("question: ",question)
        ans_str = llm.query(question)
        ans_str = ans_str[ans_str.find('{'): ans_str.rfind('}') + 1]
        # Logger.info(f'ans_str:{ans_str}')
        now_ans = json5.loads(ans_str)
        ans['problems'].extend(now_ans['problems'])
    paper_text = generate_paper_text(ans)
    # filepath = os.path.join('ai_test_system','paper')
    # os.makedirs(filepath, exist_ok=True)
    # filename = os.path.join(filepath,'test.txt')
    # with open(filename, 'w', encoding='utf-8') as f:
    #     f.write(paper_text)
    # # 创建下载文件的 Response
    # response = Response(paper_text, content_type='text/plain')
    # response.headers["Content-Disposition"] = "attachment; filename=problems.txt"
    # return response
    # 创建 Word 文档
    document = Document()
    document.add_heading('测试试卷', 0)
    document.add_paragraph(paper_text)

    # 保存文档
    filepath = os.path.join('ai_test_system', 'paper')
    os.makedirs(filepath, exist_ok=True)
    filename = os.path.join(filepath, 'test.docx')
    document.save(filename)

    # 返回 Word 文档供下载
    return send_file(filename, as_attachment=True, download_name='problems.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def generate_paper_text(problems_data):
    problems = problems_data['problems']
    
    # 创建各题型的部分列表
    single_choice_text = []
    fillin_text = []
    judgement_text = []

    option_letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J']  # 可以扩展到更多选项

    # 统计各题型数量
    problem_count = {'single_choice': 0, 'fillin': 0, 'judgement': 0}

    # 按题型分配到不同的部分
    for problem in problems:
        if problem['type'] == 'single_choice':
            problem_count['single_choice'] += 1
            single_choice_text.append(f"{problem_count['single_choice']}. {problem['problem']}")
            for i, choice in enumerate(problem['choices']):
                if i < len(option_letters):
                    single_choice_text.append(f"{option_letters[i]}. {choice}")
                else:
                    single_choice_text.append(f"{i + 1}. {choice}")  # 超过10个选项时使用数字序号

        elif problem['type'] == 'fillin':
            problem_count['fillin'] += 1
            fillin_text.append(f"{problem_count['fillin']}. {problem['problem']}")
            fillin_text.append("___________")  # 填空题用下划线表示

        elif problem['type'] == 'judgement':
            problem_count['judgement'] += 1
            judgement_text.append(f"{problem_count['judgement']}. {problem['problem']}")
            judgement_text.append("A. 是")
            judgement_text.append("B. 否")

    # 拼接各部分内容
    paper_text = []

    if single_choice_text:
        paper_text.append("一、选择题：")
        paper_text.extend(single_choice_text)

    if fillin_text:
        paper_text.append("\n二、填空题：")
        paper_text.extend(fillin_text)

    if judgement_text:
        paper_text.append("\n三、判断题：")
        paper_text.extend(judgement_text)

    # 将列表转换为字符串并返回
    return "\n".join(map(str, paper_text))

def save_to_word(filename, paper_text):
    document = Document()
    document.add_heading('试卷内容', level=1)
    document.add_paragraph(paper_text)
    document.save(filename)

@app.route('/get_historyevaluation_download_custompaper', methods=['POST'])
def get_historyevaluation_download_custompaper_handler():
    content = request.json
    Logger.info(content)
    """
    {
    "evaluation": "",
   	"shortcoming": "" ,
   	"suggestion": "" ,
   	"dimention": ["","","","","",""],
   	"score": ["90", "","","","",""],
   }
    """
    evaluation = content['evaluation']
    shortcoming = content['shortcoming']
    suggestion = content['suggestion']
    ans = {'problems': []}
    question = "请根据学生评价，为了强化学生能力出题。" + get_problems_prompt1 + "学生评价是：" + evaluation +",学生不足是：" + shortcoming + ",建议是："+ suggestion + ''
    ans_str = llm.query(question)
    ans_str = ans_str[ans_str.find('{'): ans_str.rfind('}') + 1]
    now_ans = json5.loads(ans_str)
    ans['problems'].extend(now_ans['problems'])
    paper_text = generate_paper_text(ans)
    filepath = os.path.join('ai_test_system','paper')
    os.makedirs(filepath, exist_ok=True)
    filename = os.path.join(filepath,'test.txt')
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(paper_text)
    # 创建下载文件的 Response
    response = Response(paper_text, content_type='text/plain')
    response.headers["Content-Disposition"] = "attachment; filename=problems.txt"
    return response

@app.route('/get_custom_onlinepaper', methods=['POST'])
def get_custom_onlinepaper_handler():
    content = request.json
    Logger.info(content)
    """
    {
    "evaluation": "",
   	"shortcoming": "" ,
   	"suggestion": "" ,
   	"dimention": ["","","","","",""],
   	"score": ["90", "","","","",""],
   }
    """
    evaluation = content['evaluation']
    shortcoming = content['shortcoming']
    suggestion = content['suggestion']
    ans = {'problems': []}
    question = "请根据学生评价，为了强化学生能力出题。" + get_problems_prompt1 + "学生评价是：" + evaluation +",学生不足是：" + shortcoming + ",建议是："+ suggestion + ''
    ans_str = llm.query(question)
    ans_str = ans_str[ans_str.find('{'): ans_str.rfind('}') + 1]
    now_ans = json5.loads(ans_str)
    ans['problems'].extend(now_ans['problems'])
    return jsonify(ans)  # 确保返回的是JSON格式